"""Generates a tailored .docx resume from the base resume + a TailoringPlan.

Strategy: copy the base document structure, then:
- Insert/replace a "Tailored Summary" paragraph near the top.
- Reorder the Skills section bullets in place (emphasized skills first).
We deliberately don't rewrite existing experience bullet text verbatim — only reorder
and add a clearly-labeled summary — so nothing untruthful is introduced silently.
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.resume.resume_tailor import TailoringPlan


def generate_tailored_docx(base_resume_path: Path, plan: TailoringPlan, output_path: Path, job_title: str, company_name: str) -> None:
    document = Document(str(base_resume_path))

    if plan.suggested_summary:
        _insert_summary(document, plan.suggested_summary)

    if plan.relevant_skills_section:
        _reorder_skills_section(document, plan.relevant_skills_section)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def _insert_summary(document: Document, summary_text: str) -> None:
    # Insert near the top (after the first paragraph, typically the name/header line).
    paragraphs = document.paragraphs
    if not paragraphs:
        return
    anchor = paragraphs[0]._p
    heading = document.add_paragraph()
    heading_run = heading.add_run("Tailored Summary (auto-generated — review before sending)")
    heading_run.bold = True
    heading_run.font.size = Pt(10)
    anchor.addnext(heading._p)

    body = document.add_paragraph(summary_text)
    heading._p.addnext(body._p)


def _reorder_skills_section(document: Document, ordered_skills: list[str]) -> None:
    paragraphs = document.paragraphs
    skills_idx = None
    for i, p in enumerate(paragraphs):
        if p.text.strip().lower().startswith("skill"):
            skills_idx = i
            break
    if skills_idx is None:
        return

    # Find the contiguous block of skill lines following the header and rewrite their text in place.
    i = skills_idx + 1
    skill_paragraphs = []
    while i < len(paragraphs) and paragraphs[i].text.strip():
        skill_paragraphs.append(paragraphs[i])
        i += 1

    for idx, p in enumerate(skill_paragraphs):
        if idx < len(ordered_skills):
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = ordered_skills[idx]
            else:
                p.add_run(ordered_skills[idx])
