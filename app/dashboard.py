"""BioInternshipRadar — public anonymous Streamlit dashboard.

No sign-in required. All visitor-specific data (resume, notes, tracking)
lives exclusively in st.session_state and is never written to the shared DB.
"""
# Ensure the repository root is on sys.path so `from app.*` imports work when
# Streamlit executes this file directly (e.g. `streamlit run app/dashboard.py`).
from pathlib import Path
import sys

_PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

import io
from datetime import datetime

import pandas as pd
import streamlit as st
import yaml

from app.config import settings
from app.database import get_session, init_db
from app.models import Company, Opportunity, ScanLog
from app.schemas import CompanyIn
from app.services import export_service
from app.services.company_service import (
    companies_missing_career_url,
    create_company,
    list_companies,
    sync_companies_from_yaml,
    update_company,
)
from app.services.notification_service import send_test_notification
from app.services.opportunity_service import list_opportunities, upsert_opportunity
from app.services.scan_service import run_scan_all_active, run_scan_for_company_id

STATUS_OPTIONS = [
    "New", "Reviewing", "Interested", "Applied", "Rejected", "Not Relevant", "Archived",
]

st.set_page_config(page_title="BioInternship Radar", layout="wide")
init_db()

with get_session() as _s:
    sync_companies_from_yaml(_s)


# ---------------------------------------------------------------------------
# Session state helpers — all visitor-specific data lives here only
# ---------------------------------------------------------------------------

def _tracking() -> dict:
    """Returns the per-visitor job tracking dict: {opp_id: {"status", "is_saved", "notes"}}."""
    if "job_tracking" not in st.session_state:
        st.session_state["job_tracking"] = {}
    return st.session_state["job_tracking"]


def _get_track(opp_id: int) -> dict:
    return _tracking().get(opp_id, {"status": "New", "is_saved": False, "notes": ""})


def _set_track(opp_id: int, **kwargs) -> None:
    t = _tracking()
    if opp_id not in t:
        t[opp_id] = {"status": "New", "is_saved": False, "notes": ""}
    t[opp_id].update(kwargs)


def _clear_resume_state() -> None:
    for key in [
        "resume_bytes", "resume_filename", "resume_text",
        "job_description", "tailored_result", "tailored_bytes",
        "tailoring_plan", "tailoring_opportunity_id",
    ]:
        st.session_state.pop(key, None)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    st.sidebar.title("BioInternship Radar")
    st.sidebar.caption("Public · No sign-in required")

    page = st.sidebar.radio(
        "Navigate",
        [
            "Job Feed",
            "My Tracking",
            "Overview",
            "Sources",
            "Companies",
            "Resume Tailoring",
            "Notifications",
            "Scan Logs",
            "Settings",
        ],
    )

    if page == "Job Feed":
        render_job_feed()
    elif page == "My Tracking":
        render_my_tracking()
    elif page == "Overview":
        render_overview()
    elif page == "Sources":
        render_sources()
    elif page == "Companies":
        render_companies()
    elif page == "Resume Tailoring":
        render_resume_tailoring()
    elif page == "Notifications":
        render_notifications()
    elif page == "Scan Logs":
        render_scan_logs()
    elif page == "Settings":
        render_settings()


# ---------------------------------------------------------------------------
# Page: Job Feed
# ---------------------------------------------------------------------------

def render_job_feed() -> None:
    st.header("Job Feed")

    with get_session() as session:
        companies = list_companies(session)
        opportunities = list_opportunities(session)

    last_scan = max((c.last_scanned_at for c in companies if c.last_scanned_at), default=None)

    hdr_col, btn_col = st.columns([5, 1])
    with hdr_col:
        if last_scan:
            mins = int((datetime.utcnow() - last_scan).total_seconds() / 60)
            st.caption(
                f"Last scan: {last_scan.strftime('%Y-%m-%d %H:%M UTC')} ({mins} min ago) "
                f"| {len(opportunities)} jobs in database"
            )
        else:
            st.info(
                "No scans have run yet. Click **Refresh Jobs** to discover internships, "
                "or run `python run_scanner.py --all` from the terminal."
            )
    with btn_col:
        if st.button("Refresh Jobs", type="primary", use_container_width=True):
            with st.spinner("Scanning all active sources…"):
                with get_session() as session:
                    scan_run = run_scan_all_active(session)
            st.success(
                f"Done — {scan_run.new_opportunities_found} new jobs "
                f"({scan_run.opportunities_found} total, {scan_run.errors_count} errors)"
            )
            st.rerun()

    col1, col2, col3, col4, col5 = st.columns(5)
    company_filter = col1.selectbox("Company", ["All"] + sorted({o.company_name for o in opportunities}))
    source_filter = col2.selectbox(
        "Source",
        ["All"] + sorted({o.source_platform for o in opportunities if o.source_platform not in ("", "unknown")}),
    )
    min_score = col3.slider("Min Fit Score", 0, 100, 0)
    tracking = _tracking()
    show_saved = col4.checkbox("Saved only")
    keyword = col5.text_input("Keyword search")

    filtered = opportunities
    if company_filter != "All":
        filtered = [o for o in filtered if o.company_name == company_filter]
    if source_filter != "All":
        filtered = [o for o in filtered if o.source_platform == source_filter]
    filtered = [o for o in filtered if o.fit_score >= min_score]
    if show_saved:
        filtered = [o for o in filtered if _get_track(o.id)["is_saved"]]
    if keyword:
        kw = keyword.lower()
        filtered = [
            o for o in filtered
            if kw in o.job_title.lower() or kw in o.company_name.lower() or kw in o.description.lower()
        ]

    st.caption(f"Showing {len(filtered)} of {len(opportunities)} opportunities")

    # Export shared job data — download button delivers bytes, no server-side file
    export_col_a, export_col_b = st.columns(2)
    if export_col_a.button("Export jobs to CSV"):
        with get_session() as session:
            rows = export_service._build_opportunity_rows(session)
        csv_bytes = pd.DataFrame(rows).to_csv(index=False).encode()
        export_col_a.download_button("Download CSV", csv_bytes, "opportunities.csv", "text/csv")
    if export_col_b.button("Export jobs to Excel"):
        with get_session() as session:
            rows = export_service._build_opportunity_rows(session)
        buf = io.BytesIO()
        pd.DataFrame(rows).to_excel(buf, index=False)
        export_col_b.download_button("Download Excel", buf.getvalue(), "opportunities.xlsx")

    for opp in filtered[:200]:
        track = _get_track(opp.id)
        current_status = track["status"]
        is_saved = track["is_saved"]

        with st.container(border=True):
            c1, c2 = st.columns([4, 1])
            with c1:
                score_icon = "🟢" if opp.fit_score >= 70 else ("🟡" if opp.fit_score >= 40 else "🔴")
                st.markdown(f"**{opp.job_title}** — {opp.company_name}")
                parts = []
                if opp.location:
                    parts.append(opp.location)
                if opp.posted_date:
                    parts.append(f"Posted {opp.posted_date[:10]}")
                parts.append(f"Found {opp.detected_date.strftime('%Y-%m-%d')}")
                if opp.source_platform and opp.source_platform != "unknown":
                    parts.append(f"via {opp.source_platform}")
                st.caption(" · ".join(parts))
                st.caption(f"{score_icon} Fit score: {opp.fit_score}/100")
                if opp.matched_keywords:
                    st.caption(f"Matched: {opp.matched_keywords}")
                if opp.fit_score_explanation:
                    with st.expander("Why this matched"):
                        st.text(opp.fit_score_explanation)
                if opp.job_url:
                    st.markdown(f"[Apply →]({opp.job_url})")
            with c2:
                save_label = "★ Saved" if is_saved else "☆ Save"
                if st.button(save_label, key=f"save_{opp.id}"):
                    _set_track(opp.id, is_saved=not is_saved)
                    st.rerun()
                new_status = st.selectbox(
                    "Status",
                    STATUS_OPTIONS,
                    index=STATUS_OPTIONS.index(current_status) if current_status in STATUS_OPTIONS else 0,
                    key=f"status_{opp.id}",
                )
                if new_status != current_status:
                    _set_track(opp.id, status=new_status)
                    st.rerun()

    st.divider()
    with st.expander("Add a missing job manually"):
        with st.form("manual_opportunity_form"):
            company_name = st.text_input("Company")
            job_title = st.text_input("Role Title")
            job_url = st.text_input("Job URL")
            location = st.text_input("Location")
            posted_date = st.text_input("Posted Date (optional)")
            description = st.text_area("Description / paste job posting text")
            submitted = st.form_submit_button("Add")
            if submitted and company_name and job_title:
                from app.schemas import OpportunityCandidate
                with get_session() as session:
                    matching_company = next(
                        (c for c in list_companies(session) if c.name.lower() == company_name.lower()), None
                    )
                    candidate = OpportunityCandidate(
                        company_name=company_name, job_title=job_title, job_url=job_url,
                        location=location, posted_date=posted_date, description=description,
                        source_platform="manual",
                    )
                    opp, is_new = upsert_opportunity(session, candidate, matching_company)
                st.success(
                    f"{'Added' if is_new else 'Already existed'}: "
                    f"{job_title} at {company_name} — fit score {opp.fit_score}"
                )


# ---------------------------------------------------------------------------
# Page: My Tracking  (temporary — session only)
# ---------------------------------------------------------------------------

def render_my_tracking() -> None:
    st.header("My Tracking")
    st.info(
        "**Temporary for this browser session.** "
        "Saved jobs and statuses are stored only in this browser tab — "
        "they will be lost when you close or refresh the page. "
        "Export your data before leaving."
    )
    st.warning("Refreshing or closing this page will erase your temporary tracking data.")

    tracking = _tracking()
    if not tracking:
        st.caption("No jobs saved or tracked yet. Use the Job Feed to save jobs or update their status.")
        return

    tracked_ids = list(tracking.keys())
    with get_session() as session:
        opportunities = {o.id: o for o in list_opportunities(session) if o.id in tracked_ids}

    status_filter = st.selectbox("Filter by status", ["All"] + STATUS_OPTIONS)

    rows_shown = 0
    for opp_id, track in tracking.items():
        if status_filter != "All" and track["status"] != status_filter:
            continue
        opp = opportunities.get(opp_id)
        if not opp:
            continue
        rows_shown += 1
        with st.container(border=True):
            c1, c2 = st.columns([4, 1])
            with c1:
                score_icon = "🟢" if opp.fit_score >= 70 else ("🟡" if opp.fit_score >= 40 else "🔴")
                st.markdown(f"**{opp.job_title}** — {opp.company_name}")
                st.caption(f"Status: {track['status']} · {score_icon} Fit: {opp.fit_score}")
                if opp.location:
                    st.caption(opp.location)
                if track["notes"]:
                    st.caption(f"Notes: {track['notes']}")
                if opp.job_url:
                    st.markdown(f"[Open →]({opp.job_url})")
            with c2:
                new_status = st.selectbox(
                    "Status",
                    STATUS_OPTIONS,
                    index=STATUS_OPTIONS.index(track["status"]) if track["status"] in STATUS_OPTIONS else 0,
                    key=f"mytrack_status_{opp_id}",
                )
                if new_status != track["status"]:
                    _set_track(opp_id, status=new_status)
                    st.rerun()

        with st.expander(f"Notes — {opp.job_title}"):
            notes_val = st.text_area("Notes", track["notes"], key=f"mytrack_notes_{opp_id}")
            if st.button("Save notes", key=f"mytrack_savenotes_{opp_id}"):
                _set_track(opp_id, notes=notes_val)
                st.success("Saved (temporary — session only)")

    st.caption(f"Showing {rows_shown} tracked jobs")

    st.divider()
    # Export personal tracking data — delivered as download, never written to shared DB
    if st.button("Export my tracking to CSV"):
        with get_session() as session:
            opps = {o.id: o for o in list_opportunities(session)}
        export_rows = []
        for opp_id, track in tracking.items():
            opp = opps.get(opp_id)
            if not opp:
                continue
            export_rows.append({
                "job_title": opp.job_title,
                "company_name": opp.company_name,
                "job_url": opp.job_url,
                "fit_score": opp.fit_score,
                "status": track["status"],
                "is_saved": track["is_saved"],
                "notes": track["notes"],
            })
        csv_bytes = pd.DataFrame(export_rows).to_csv(index=False).encode()
        st.download_button("Download CSV", csv_bytes, "my_tracking.csv", "text/csv")


# ---------------------------------------------------------------------------
# Page: Overview
# ---------------------------------------------------------------------------

def render_overview() -> None:
    st.header("Overview")

    with get_session() as session:
        companies = list_companies(session)
        opportunities = list_opportunities(session)
        missing_url = companies_missing_career_url(session)

    tracking = _tracking()
    today = datetime.utcnow().date()
    new_today = [o for o in opportunities if o.detected_date.date() == today]
    high_fit = [o for o in opportunities if o.fit_score >= settings.min_notification_fit_score]
    last_scan = max((c.last_scanned_at for c in companies if c.last_scanned_at), default=None)
    my_saved = [opp_id for opp_id, t in tracking.items() if t["is_saved"]]
    my_applied = [opp_id for opp_id, t in tracking.items() if t["status"] == "Applied"]

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Total Jobs", len(opportunities))
    col2.metric("New Today", len(new_today))
    col3.metric("High-Fit (≥60)", len(high_fit))
    col4.metric("Last Scan", last_scan.strftime("%Y-%m-%d %H:%M") if last_scan else "Never")

    col5, col6, col7, col8 = st.columns(4)
    col5.metric("My Saved (session)", len(my_saved))
    col6.metric("My Applied (session)", len(my_applied))
    col7.metric("Companies Tracked", len(companies))
    col8.metric("Need Config", len(missing_url))

    if not opportunities:
        st.info("No opportunities yet. Click **Job Feed → Refresh Jobs** to start discovery.")
        return

    df = pd.DataFrame([_opportunity_to_row(o) for o in opportunities])
    company_to_category = {c.name: c.category for c in companies}
    df["category"] = df["company_name"].map(company_to_category).fillna("Uncategorized")

    col_l, col_r = st.columns(2)
    with col_l:
        st.subheader("Jobs by Category")
        st.bar_chart(df["category"].value_counts())
    with col_r:
        st.subheader("Jobs by Company (top 15)")
        st.bar_chart(df["company_name"].value_counts().head(15))

    st.subheader("New Jobs Over Time")
    df["detected_day"] = pd.to_datetime(df["detected_date"]).dt.date
    st.line_chart(df.groupby("detected_day").size())


# ---------------------------------------------------------------------------
# Page: Sources
# ---------------------------------------------------------------------------

def render_sources() -> None:
    st.header("Sources")
    st.caption("Which employers are configured and what each provider retrieved last scan.")

    with get_session() as session:
        companies = list_companies(session)

    scannable = [c for c in companies if c.platform not in ("workday", "icims", "unknown", "company_static") or getattr(c, "board_id", "")]
    manual = [c for c in companies if c not in scannable]

    col1, col2, col3 = st.columns(3)
    active_scannable = [c for c in scannable if c.active]
    col1.metric("Auto-scannable sources", len(active_scannable))
    col2.metric("Manual-only sources", len(manual))
    errors = [c for c in companies if c.last_scan_status == "error"]
    col3.metric("Sources with errors", len(errors))

    st.subheader("Auto-scannable Companies")
    if active_scannable:
        rows = []
        for c in active_scannable:
            with get_session() as session:
                opp_count = session.query(Opportunity).filter(Opportunity.company_name == c.name).count()
            rows.append({
                "Company": c.name,
                "Provider": c.platform,
                "Board ID": getattr(c, "board_id", "") or "",
                "Enabled": c.active,
                "Last Scan": c.last_scanned_at.strftime("%Y-%m-%d %H:%M") if c.last_scanned_at else "Never",
                "Status": c.last_scan_status,
                "Jobs Found": opp_count,
                "Error": c.last_scan_error[:80] if c.last_scan_error else "",
            })
        st.dataframe(pd.DataFrame(rows), hide_index=True, use_container_width=True)
    else:
        st.info("No auto-scannable companies configured yet.")

    st.subheader("Manual-only Companies (Workday / Unknown ATS)")
    st.caption("These companies require you to visit their career page directly.")
    if manual:
        rows = [{"Company": c.name, "Category": c.category, "Platform": c.platform, "Notes": c.notes[:100]} for c in manual]
        st.dataframe(pd.DataFrame(rows), hide_index=True, use_container_width=True)

    st.divider()
    st.subheader("Add a New Scannable Source")
    st.markdown("""
To add a company:
1. Find its ATS: visit `boards.greenhouse.io/<slug>`, `jobs.lever.co/<slug>`, or `jobs.ashbyhq.com/<slug>`
2. If the job board loads, add to `data/companies.yaml`:
   ```yaml
   - name: Example Biotech
     platform: greenhouse
     board_id: examplebiotech
     active: true
   ```
3. Click **Refresh Jobs** on the Job Feed page.
    """)


# ---------------------------------------------------------------------------
# Page: Companies
# ---------------------------------------------------------------------------

def render_companies() -> None:
    st.header("Companies")

    with st.expander("Add a new company"):
        with st.form("add_company_form"):
            name = st.text_input("Name")
            category = st.selectbox("Category", ["Bio Companies", "Startups / Healthcare Tech", "Consulting Companies"])
            career_url = st.text_input("Career URL (optional if board_id is set)")
            internship_url = st.text_input("Internship URL (optional)")
            platform = st.selectbox("Platform", ["unknown", "greenhouse", "lever", "ashby", "workday", "icims", "company_static"])
            board_id = st.text_input("Board ID (Greenhouse token / Lever slug / Ashby slug)")
            priority = st.selectbox("Priority", ["High", "Medium", "Low"], index=1)
            submitted = st.form_submit_button("Add Company")
            if submitted and name:
                with get_session() as session:
                    create_company(session, CompanyIn(
                        name=name, category=category, career_url=career_url, internship_url=internship_url,
                        platform=platform, board_id=board_id, priority=priority,
                    ))
                st.success(f"Added {name}")
                st.rerun()

    col_a, col_b, col_c = st.columns(3)
    with col_a:
        if st.button("Export companies to CSV"):
            with get_session() as session:
                path = export_service.export_companies_csv(session)
            st.success(f"Exported to {path}")
    with col_b:
        if st.button("Export companies to YAML"):
            with get_session() as session:
                path = export_service.export_companies_yaml(session)
            st.success(f"Exported to {path}")
    with col_c:
        if st.button("Run Scan Now (all active companies)"):
            with get_session() as session:
                scan_run = run_scan_all_active(session)
            st.success(f"Scan complete: {scan_run.new_opportunities_found} new opportunities")

    with get_session() as session:
        companies = list_companies(session)
        rows = [_company_to_row(c) for c in companies]

    if not rows:
        st.info("No companies loaded yet.")
        return

    df = pd.DataFrame(rows)
    not_scannable = df[(df["career_url"] == "") & (df["board_id"] == "") & (~df["platform"].isin(["workday"]))]
    if len(not_scannable):
        st.warning(
            f"{len(not_scannable)} companies have no scan config (no career_url, no board_id). "
            "Set platform + board_id in companies.yaml, or add career_url."
        )

    st.dataframe(df, use_container_width=True, hide_index=True)

    st.subheader("Edit a company")
    with get_session() as session:
        companies = list_companies(session)
    selected_name = st.selectbox("Select company to edit", [c.name for c in companies])
    selected = next((c for c in companies if c.name == selected_name), None)
    if selected:
        with st.form("edit_company_form"):
            career_url = st.text_input("Career URL", value=selected.career_url)
            internship_url = st.text_input("Internship URL", value=selected.internship_url)
            _platforms = ["unknown", "greenhouse", "lever", "ashby", "workday", "icims", "company_static"]
            _plat_idx = _platforms.index(selected.platform) if selected.platform in _platforms else 0
            platform = st.selectbox("Platform", _platforms, index=_plat_idx)
            board_id = st.text_input("Board ID", value=getattr(selected, "board_id", "") or "")
            priority = st.selectbox("Priority", ["High", "Medium", "Low"], index=["High", "Medium", "Low"].index(selected.priority or "Medium"))
            active = st.checkbox("Active", value=selected.active)
            notes = st.text_area("Notes", value=selected.notes)
            col1, col2 = st.columns(2)
            save = col1.form_submit_button("Save")
            test_scan = col2.form_submit_button("Test Scan This Company")

            if save:
                with get_session() as session:
                    update_company(session, selected.id, {
                        "career_url": career_url, "internship_url": internship_url, "platform": platform,
                        "board_id": board_id, "priority": priority, "active": active, "notes": notes,
                    })
                st.success("Saved")
                st.rerun()

            if test_scan:
                with get_session() as session:
                    scan_run = run_scan_for_company_id(session, selected.id)
                st.success(f"Test scan: {scan_run.opportunities_found} jobs found, {scan_run.new_opportunities_found} new")


# ---------------------------------------------------------------------------
# Page: Resume Tailoring
# ---------------------------------------------------------------------------

def _extract_text_from_upload(uploaded_file) -> str:
    """Extract plain text from an uploaded PDF or DOCX file — entirely in memory."""
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    elif name.endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    return raw.decode("utf-8", errors="replace")


def render_resume_tailoring() -> None:
    st.header("Resume Tailoring")

    # Privacy notice
    st.info(
        "**Privacy:** Your resume and generated results are kept only for your current browser session. "
        "They are not saved to the shared job database. "
        "Download your result before closing or refreshing the page."
    )
    st.warning("Refreshing or closing this page may erase your temporary work.")

    # Delete / Start over button (always visible when there is session data)
    has_data = any(
        k in st.session_state
        for k in ["resume_bytes", "resume_text", "tailored_result", "tailored_bytes"]
    )
    if has_data:
        if st.button("Delete my resume and results", type="secondary"):
            _clear_resume_state()
            st.success("Your resume and all generated results have been deleted from this session.")
            st.rerun()

    st.divider()

    # Step 1: Upload resume
    st.subheader("Step 1 — Upload your resume")
    uploaded = st.file_uploader(
        "Upload PDF or DOCX",
        type=["pdf", "docx"],
        key="resume_uploader",
        help="Your file is read in memory only and never saved to disk.",
    )

    if uploaded is not None:
        new_bytes = uploaded.read()
        uploaded.seek(0)
        # Auto-clear previous results when a new resume is uploaded
        if st.session_state.get("resume_filename") != uploaded.name:
            _clear_resume_state()
        st.session_state["resume_bytes"] = new_bytes
        st.session_state["resume_filename"] = uploaded.name
        try:
            st.session_state["resume_text"] = _extract_text_from_upload(uploaded)
            st.caption(f"Loaded: **{uploaded.name}** ({len(new_bytes):,} bytes, {len(st.session_state['resume_text'])} chars extracted)")
        except Exception as exc:
            st.error(f"Could not read resume: {exc}")
            return

    resume_text = st.session_state.get("resume_text", "")
    if not resume_text:
        st.caption("Upload a resume above to continue.")
        return

    # Step 2: Select or paste job description
    st.divider()
    st.subheader("Step 2 — Select a job or paste a description")

    with get_session() as session:
        opportunities = list_opportunities(session)

    source_mode = st.radio("Job description source", ["Pick from job feed", "Paste manually"], horizontal=True)

    if source_mode == "Pick from job feed":
        if not opportunities:
            st.info("No jobs in the database yet. Run a scan first.")
            return
        labels = [f"{o.company_name} — {o.job_title}" for o in opportunities]
        selected_idx = st.selectbox("Select opportunity", range(len(opportunities)), format_func=lambda i: labels[i])
        opp = opportunities[selected_idx]
        job_description = opp.description
        job_label = f"{opp.company_name} — {opp.job_title}"
    else:
        job_label = st.text_input("Job title / company (for labeling)")
        job_description = st.text_area("Paste job description here", height=200)
        opp = None

    if not job_description.strip():
        st.caption("Provide a job description to continue.")
        return

    # Step 3: Generate
    st.divider()
    st.subheader("Step 3 — Generate tailored content")

    use_llm = st.checkbox(
        f"Use LLM-assisted tailoring (provider: {settings.llm_provider})",
        value=False,
        disabled=(settings.llm_provider == "none"),
    )
    if use_llm:
        st.warning("LLM tailoring sends your resume text and the job description to the configured LLM provider.")

    if st.button("Generate tailored resume", type="primary"):
        _clear_resume_state()
        st.session_state["resume_text"] = resume_text  # restore after clear

        with st.spinner("Generating…"):
            try:
                if opp is not None and use_llm:
                    # Use existing resume_service with the shared opp object for plan generation
                    from app.services.resume_service import preview_tailoring
                    plan, _ = preview_tailoring(opp, use_llm=use_llm)
                    tailored_text = _build_tailored_text(resume_text, job_description, plan)
                    st.session_state["tailoring_plan"] = plan
                else:
                    from app.services.keyword_service import load_keywords
                    kw_dict = load_keywords()
                    # Flatten all keyword lists into a single list
                    keywords = [kw for group in kw_dict.values() for kw in group]
                    tailored_text = _build_tailored_text_simple(resume_text, job_description, keywords)

                st.session_state["tailored_result"] = tailored_text
                st.session_state["job_description"] = job_description

                # Build a downloadable DOCX in memory
                from docx import Document as DocxDocument
                doc = DocxDocument()
                doc.add_heading(job_label or "Tailored Resume", 0)
                for line in tailored_text.split("\n"):
                    doc.add_paragraph(line)
                buf = io.BytesIO()
                doc.save(buf)
                st.session_state["tailored_bytes"] = buf.getvalue()
            except Exception as exc:
                st.error(f"Generation failed: {exc}")
                return

    tailored = st.session_state.get("tailored_result")
    tailored_bytes = st.session_state.get("tailored_bytes")

    if tailored:
        st.divider()
        st.subheader("Result")
        st.text_area("Tailored content (read-only)", tailored, height=400, disabled=True)

        dl_col, del_col = st.columns(2)
        if tailored_bytes:
            dl_col.download_button(
                "Download tailored resume (.docx)",
                tailored_bytes,
                file_name=f"tailored_resume_{(job_label or 'resume').replace(' ', '_')[:40]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        if del_col.button("Delete my resume and results", key="delete_bottom"):
            _clear_resume_state()
            st.success("Deleted.")
            st.rerun()

        st.divider()
        st.subheader("Outreach Drafts")
        from app.resume.cover_letter_generator import (
            generate_linkedin_message,
            generate_networking_message,
            generate_referral_request,
        )
        from app.services.resume_service import load_student_profile
        profile = load_student_profile()
        student_name = profile.get("student_name", "")
        company = opp.company_name if opp else (job_label or "the company")
        title = opp.job_title if opp else (job_label or "Internship")
        plan = st.session_state.get("tailoring_plan")

        col1, col2, col3 = st.columns(3)
        with col1:
            st.text_area("LinkedIn message", generate_linkedin_message(student_name, company, title), height=120)
        with col2:
            st.text_area("Networking email", generate_networking_message(student_name, company, title, plan), height=120)
        with col3:
            st.text_area("Referral request", generate_referral_request(student_name, company, title), height=120)


def _build_tailored_text(resume_text: str, job_description: str, plan) -> str:
    """Build a simple tailored text block from plan output."""
    lines = [
        "=== TAILORED RESUME GUIDANCE ===",
        "",
        "Suggested Summary:",
        plan.suggested_summary or "(no summary generated)",
        "",
        "Matched Keywords:",
        ", ".join(plan.matched_keywords) if plan.matched_keywords else "None",
        "",
        "Missing Keywords (do not fabricate — for awareness only):",
        ", ".join(plan.missing_keywords) if plan.missing_keywords else "None",
        "",
        "Changes to Apply:",
    ]
    for change in (plan.changes_made or []):
        lines.append(f"- {change}")
    lines += ["", "=== ORIGINAL RESUME TEXT ===", "", resume_text]
    return "\n".join(lines)


def _build_tailored_text_simple(resume_text: str, job_description: str, keywords: list) -> str:
    """Keyword-based tailoring without LLM."""
    jd_lower = job_description.lower()
    matched = [kw for kw in keywords if kw.lower() in jd_lower]
    lines = [
        "=== TAILORING NOTES (keyword match, no LLM) ===",
        "",
        f"Keywords found in job description ({len(matched)}):",
        ", ".join(matched) if matched else "None found",
        "",
        "Review your resume and emphasise the above keywords where genuinely applicable.",
        "",
        "=== ORIGINAL RESUME TEXT ===",
        "",
        resume_text,
    ]
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Page: Notifications
# ---------------------------------------------------------------------------

def render_notifications() -> None:
    st.header("Notifications")

    col1, col2, col3 = st.columns(3)
    with col1:
        st.subheader("Email")
        st.write(f"Enabled: {settings.email_enabled}")
        if st.button("Send test email"):
            success, error = send_test_notification("email")
            st.success("Sent") if success else st.error(error)
    with col2:
        st.subheader("Telegram")
        st.write(f"Enabled: {settings.telegram_enabled}")
        if st.button("Send test Telegram"):
            success, error = send_test_notification("telegram")
            st.success("Sent") if success else st.error(error)
    with col3:
        st.subheader("Slack")
        st.write(f"Enabled: {settings.slack_enabled}")
        if st.button("Send test Slack"):
            success, error = send_test_notification("slack")
            st.success("Sent") if success else st.error(error)

    st.caption("Configure channels in .env (see .env.example). Restart after changing .env.")


# ---------------------------------------------------------------------------
# Page: Scan Logs
# ---------------------------------------------------------------------------

def render_scan_logs() -> None:
    st.header("Scan Logs")

    with get_session() as session:
        logs = session.query(ScanLog).order_by(ScanLog.started_at.desc()).limit(300).all()

    if not logs:
        st.info("No scans have run yet.")
        return

    df = pd.DataFrame([{
        "started_at": l.started_at,
        "company": l.company_name,
        "scanner": l.scanner_used,
        "status": l.status,
        "jobs_found": l.jobs_found,
        "new_jobs": l.new_jobs_found,
        "error": l.error_message,
    } for l in logs])
    st.dataframe(df, use_container_width=True, hide_index=True)

    if st.button("Export scan logs to CSV"):
        with get_session() as session:
            path = export_service.export_scan_logs_csv(session)
        st.success(f"Exported to {path}")


# ---------------------------------------------------------------------------
# Page: Settings
# ---------------------------------------------------------------------------

def render_settings() -> None:
    st.header("Settings")

    st.subheader("Scan Configuration (from .env)")
    st.code(
        f"SCAN_INTERVAL_MINUTES={settings.scan_interval_minutes}\n"
        f"INTERNSHIP_ONLY={settings.internship_only}\n"
        f"MIN_NOTIFICATION_FIT_SCORE={settings.min_notification_fit_score}\n"
        f"LLM_PROVIDER={settings.llm_provider}\n"
    )

    for label, path_attr in [
        ("Keyword Lists (data/keywords.yaml)", "keywords_yaml_path"),
        ("Ignored Keywords (data/ignored_keywords.yaml)", "ignored_keywords_yaml_path"),
        ("Student Profile (data/student_profile.yaml)", "student_profile_yaml_path"),
    ]:
        yaml_path = getattr(settings, path_attr)
        if yaml_path.exists():
            st.subheader(label)
            with open(yaml_path, "r", encoding="utf-8") as f:
                content = f.read()
            edited = st.text_area(label, content, height=200, key=path_attr)
            if st.button(f"Save {yaml_path.name}", key=f"save_{path_attr}"):
                try:
                    yaml.safe_load(edited)
                except yaml.YAMLError as exc:
                    st.error(f"Invalid YAML: {exc}")
                else:
                    with open(yaml_path, "w", encoding="utf-8") as f:
                        f.write(edited)
                    st.success("Saved.")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _opportunity_to_row(o: Opportunity) -> dict:
    return {
        "id": o.id,
        "company_name": o.company_name,
        "job_title": o.job_title,
        "fit_score": o.fit_score,
        "detected_date": o.detected_date,
        "location": o.location,
        "source_platform": o.source_platform,
    }


def _company_to_row(c: Company) -> dict:
    return {
        "name": c.name,
        "category": c.category,
        "platform": c.platform,
        "board_id": getattr(c, "board_id", "") or "",
        "career_url": c.career_url,
        "priority": c.priority,
        "active": c.active,
        "last_scanned_at": c.last_scanned_at,
        "last_scan_status": c.last_scan_status,
    }


if __name__ == "__main__":
    main()
else:
    main()
