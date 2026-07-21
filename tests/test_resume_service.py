from pathlib import Path

from docx import Document

from app.resume.resume_parser import find_section, parse_resume
from app.resume.resume_tailor import build_tailoring_plan
from app.utils.file_utils import generated_metadata_path, generated_resume_path, slugify


def _build_sample_resume(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Jordan Student")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python")
    doc.add_paragraph("Cell Culture")
    doc.add_paragraph("Data Analysis")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Research Intern: Supported cell culture experiments for vaccine development.")
    doc.save(str(path))


def test_parse_resume_extracts_sections(tmp_path):
    resume_path = tmp_path / "base_resume.docx"
    _build_sample_resume(resume_path)

    parsed = parse_resume(resume_path)
    skills_section = find_section(parsed, "skill")
    assert skills_section is not None
    assert "Python" in skills_section.lines


def test_build_tailoring_plan_only_uses_existing_skills(tmp_path):
    resume_path = tmp_path / "base_resume.docx"
    _build_sample_resume(resume_path)
    parsed = parse_resume(resume_path)

    plan = build_tailoring_plan(
        parsed,
        job_title="Process Development Intern",
        job_description="Looking for an intern with cell culture and regulatory experience.",
        company_name="Regeneron",
    )

    assert "cell culture" in plan.matched_keywords
    assert "regulatory" in plan.missing_keywords  # not on resume — must not be fabricated
    assert all(skill in ["Python", "Cell Culture", "Data Analysis"] for skill in plan.relevant_skills_section)


def test_generated_resume_path_format():
    path = generated_resume_path("Regeneron", "Process Development Intern")
    assert path.name.endswith("_regeneron_process_development_intern_resume.docx")
    assert slugify("Regeneron") in str(path)


def test_generated_metadata_path_matches_resume_filename():
    resume_path = generated_resume_path("Regeneron", "Process Development Intern")
    metadata_path = generated_metadata_path(resume_path)
    assert metadata_path.suffix == ".json"
    assert "metadata" in metadata_path.name
